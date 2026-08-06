"""Plain text, Markdown, and DOCX.

These formats have no page geometry, so their spans carry a line range instead of a
bounding box. The eval's tolerance for them is "section heading or line range overlapping
the excerpt, ±3 lines", which is why line numbers are tracked exactly rather than
approximated from character offsets.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from reef.extract.base import (
    ExtractionError,
    ExtractionResult,
    RawPage,
    RawSpan,
    SpanKind,
    TextSource,
)
from reef.extract.tabular import _decode

_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_MD_TABLE_ROW = re.compile(r"^\s*\|.*\|\s*$")
_MD_TABLE_DIVIDER = re.compile(r"^\s*\|[\s:|-]+\|\s*$")
_LIST_ITEM = re.compile(r"^\s*([-*+]|\d+\.)\s+")


def _paragraph_blocks(lines: list[str]) -> list[tuple[int, int, list[str]]]:
    """Group lines into blocks separated by blank lines.

    Returns (first_line, last_line, lines) with 1-indexed line numbers, because the anchor
    a human checks is a line number in their editor, not an offset.
    """
    blocks: list[tuple[int, int, list[str]]] = []
    current: list[str] = []
    start = 1

    for number, line in enumerate(lines, start=1):
        if line.strip():
            if not current:
                start = number
            current.append(line)
        elif current:
            blocks.append((start, number - 1, current))
            current = []

    if current:
        blocks.append((start, len(lines), current))
    return blocks


def extract_text(path: Path | None, data: bytes, markdown: bool = False) -> ExtractionResult:
    text, encoding = _decode(data)
    lines = text.splitlines()
    if not any(line.strip() for line in lines):
        raise ExtractionError("file is empty")

    result = ExtractionResult(text_source=TextSource.NATIVE)
    result.pages.append(RawPage(number=1, width=0.0, height=0.0))

    cursor = 0
    # Heading stack, so a paragraph deep in a document knows the section path above it.
    # "Master Lease › §12 Assignment ›" is what makes a retrieved chunk interpretable
    # without its neighbours.
    heading_stack: list[tuple[int, str]] = []

    for first_line, last_line, block_lines in _paragraph_blocks(lines):
        block_text = "\n".join(block_lines).strip()
        if not block_text:
            continue

        kind = SpanKind.TEXT
        breadcrumb = " › ".join(title for _, title in heading_stack)

        if markdown:
            heading = _MD_HEADING.match(block_lines[0])
            if heading and len(block_lines) == 1:
                level = len(heading.group(1))
                title = heading.group(2).strip()
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                # Breadcrumb for a heading is the path above it, not including itself.
                breadcrumb = " › ".join(t for _, t in heading_stack)
                heading_stack.append((level, title))
                kind = SpanKind.HEADING
            elif all(_MD_TABLE_ROW.match(ln) for ln in block_lines):
                kind = SpanKind.TABLE
            elif _LIST_ITEM.match(block_lines[0]):
                kind = SpanKind.LIST_ITEM

        locator = (
            f"line {first_line}" if first_line == last_line else f"lines {first_line}-{last_line}"
        )
        result.spans.append(
            RawSpan(
                page_number=1,
                char_start=cursor,
                char_end=cursor + len(block_text),
                text=block_text,
                locator=locator,
                kind=kind,
                breadcrumb=breadcrumb,
            )
        )
        cursor += len(block_text) + 1

    result.extra = {"encoding": encoding, "line_count": len(lines), "page_count": 1}
    return result


def extract_markdown(path: Path | None, data: bytes) -> ExtractionResult:
    result = extract_text(path, data, markdown=True)
    # Markdown tables are split into per-row spans as well as the whole-table span, since
    # a question about one supplier should retrieve that supplier's row rather than a
    # forty-row table the model then has to re-read.
    expanded: list[RawSpan] = []
    for span in result.spans:
        expanded.append(span)
        if span.kind is not SpanKind.TABLE:
            continue
        base_line = _first_line_of(span.locator)
        offset = span.char_start
        for index, line in enumerate(span.text.splitlines()):
            if _MD_TABLE_DIVIDER.match(line):
                offset += len(line) + 1
                continue
            cleaned = line.strip().strip("|").strip()
            if cleaned:
                expanded.append(
                    RawSpan(
                        page_number=span.page_number,
                        char_start=offset,
                        char_end=offset + len(line),
                        text=cleaned,
                        locator=f"line {base_line + index}",
                        kind=SpanKind.TABLE_ROW,
                        breadcrumb=span.breadcrumb,
                    )
                )
            offset += len(line) + 1
    result.spans = expanded
    return result


def _first_line_of(locator: str) -> int:
    match = re.search(r"lines? (\d+)", locator)
    return int(match.group(1)) if match else 1


def extract_docx(path: Path | None, data: bytes) -> ExtractionResult:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("python-docx is not installed") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"document could not be opened: {exc}") from exc

    result = ExtractionResult(text_source=TextSource.NATIVE)
    result.pages.append(RawPage(number=1, width=0.0, height=0.0))
    cursor = 0
    ordinal = 0
    heading_stack: list[tuple[int, str]] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        ordinal += 1

        style = (paragraph.style.name or "") if paragraph.style else ""
        kind = SpanKind.TEXT
        breadcrumb = " › ".join(title for _, title in heading_stack)

        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
            except (ValueError, IndexError):
                level = 1
            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            breadcrumb = " › ".join(t for _, t in heading_stack)
            heading_stack.append((level, text))
            kind = SpanKind.HEADING
        elif style.startswith("List"):
            kind = SpanKind.LIST_ITEM

        result.spans.append(
            RawSpan(
                page_number=1,
                char_start=cursor,
                char_end=cursor + len(text),
                text=text,
                # DOCX has no fixed pagination — page breaks depend on the renderer — so
                # claiming a page number would be inventing an anchor that cannot be
                # verified. Paragraph ordinal is stable and checkable.
                locator=f"paragraph {ordinal}",
                kind=kind,
                breadcrumb=breadcrumb,
            )
        )
        cursor += len(text) + 1

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue
            text = "\t".join(cells)
            result.spans.append(
                RawSpan(
                    page_number=1,
                    char_start=cursor,
                    char_end=cursor + len(text),
                    text=text,
                    locator=f"table {table_index} row {row_index}",
                    kind=SpanKind.TABLE_ROW,
                )
            )
            cursor += len(text) + 1

    if not result.spans:
        raise ExtractionError("document contains no text")

    result.extra = {
        "paragraph_count": ordinal,
        "table_count": len(document.tables),
        "page_count": 1,
    }
    return result
